from docx import Document
from pathlib import Path

doc = Document(r"c:\Users\restd\Downloads\Content.docx")
out = Path(r"c:\Users\restd\astreanox-new\coe_extract.txt")
lines = [f"PARAS: {len(doc.paragraphs)}", f"TABLES: {len(doc.tables)}", ""]
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip():
        lines.append(f"[{i}] {t}")
lines.append("\n---TABLES---\n")
for ti, table in enumerate(doc.tables):
    lines.append(f"TABLE {ti}:")
    for row in table.rows:
        lines.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
    lines.append("")
out.write_text("\n".join(lines), encoding="utf-8")
print("WROTE", out)
